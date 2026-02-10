"""
Enhanced Resume Parser Module
Extracts comprehensive data from PDF and DOCX resumes.
"""

import re
from typing import Dict, List, Any, Optional
from io import BytesIO
from datetime import datetime

# PDF parsing
from pypdf import PdfReader

# DOCX parsing
from docx import Document

# For AI-powered extraction
from openai import OpenAI
from app.config import settings

# Use OPENAI_API_KEY for resume parsing (not APP_API_KEY)
client = OpenAI(api_key=settings.OPENAI_API_KEY) if settings.OPENAI_API_KEY else None


class ResumeParser:
    """Parse resume files (PDF/DOCX) and extract comprehensive structured data."""
    
    def __init__(self):
        self.client = client
    
    def parse_file(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Parse a resume file and extract comprehensive structured data."""
        # Extract text based on file type
        if filename.lower().endswith('.pdf'):
            text = self._extract_pdf_text(content)
        elif filename.lower().endswith('.docx'):
            text = self._extract_docx_text(content)
        else:
            raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")
        
        # Parse using comprehensive extraction
        parsed_data = self._parse_resume_comprehensive(text)
        
        return {
            "raw_text": text,
            "parsed_data": parsed_data,
            "ats_score": self._calculate_comprehensive_ats_score(parsed_data),
        }
    
    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF file with better formatting preservation."""
        try:
            pdf_file = BytesIO(content)
            reader = PdfReader(pdf_file)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n\n".join(text_parts).strip()
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX file with structure preservation."""
        try:
            docx_file = BytesIO(content)
            doc = Document(docx_file)
            text_parts = []
            
            # Extract paragraphs with formatting hints
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract table content (often contains work experience, education)
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts).strip()
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    def _parse_resume_comprehensive(self, text: str) -> Dict[str, Any]:
        """Comprehensive resume parsing with multiple extraction methods."""
        # Try AI-powered extraction first (more comprehensive)
        if self.client:
            try:
                ai_parsed = self._ai_comprehensive_extraction(text)
                if ai_parsed:
                    return ai_parsed
            except Exception as e:
                print(f"AI extraction failed: {e}")
        
        # Fallback to enhanced rule-based extraction
        return self._enhanced_rule_based_extraction(text)
    
    def _ai_comprehensive_extraction(self, text: str) -> Dict[str, Any]:
        """Use OpenAI to extract all resume sections comprehensively."""
        if not self.client:
            return {}
        
        # Handle large resumes by taking more content
        max_chars = 25000  # Increased from 8000
        text_to_parse = text[:max_chars] if len(text) > max_chars else text
        
        prompt = f"""You are an expert resume parser. Extract ALL information from this resume and return as a detailed JSON object.

CRITICAL INSTRUCTIONS:
1. Extract EVERY work experience entry found in the resume - do not skip any
2. For each work experience, capture:
   - Full job title (including seniority level like "Sr", "Lead", "Senior")
   - Complete company name
   - Employment dates (start and end)
   - Full description with all responsibilities
   - All achievements/bullet points
3. Parse dates in format: "Jan 2020", "2020", or "2024-06 - Current"
4. For current positions, set "is_current": true and "end_date": ""
5. Extract ALL projects with their URLs, tech stacks, and descriptions
6. Extract ALL skills mentioned anywhere in the resume

Return this exact structure:
{{
  "full_name": "extracted full name",
  "email": "email address",
  "phone": "phone number",
  "location": "city, state/country",
  "linkedin_url": "LinkedIn URL if present",
  "portfolio_url": "portfolio/personal website if present",
  "github_url": "GitHub URL if present",
  "summary": "professional summary/objective",
  "skills": ["skill1", "skill2", ...],
  "work_experience": [
    {{
      "title": "Complete Job Title",
      "company": "Full Company Name",
      "location": "City, State/Country",
      "start_date": "MM/YYYY or YYYY",
      "end_date": "MM/YYYY or YYYY or empty string if current",
      "is_current": true/false,
      "description": "Complete job description and all responsibilities",
      "achievements": ["achievement1", "achievement2", "achievement3", ...]
    }}
  ],
  "education": [
    {{
      "degree": "Full Degree Name (e.g., Bachelor of Technology: Electronics & Communication)",
      "institution": "University/School Name",
      "location": "City, State/Country",
      "graduation_date": "YYYY",
      "gpa": "GPA if mentioned"
    }}
  ],
  "certifications": [
    {{
      "name": "Certification Name",
      "issuer": "Issuing Organization",
      "date": "Date obtained"
    }}
  ],
  "projects": [
    {{
      "name": "Project Name",
      "description": "Detailed project description",
      "technologies": ["tech1", "tech2", "tech3"],
      "url": "Project URL if any"
    }}
  ],
  "languages": [
    {{
      "language": "Language name",
      "proficiency": "Native/Fluent/Professional/Conversational/Basic"
    }}
  ]
}}

IMPORTANT RULES:
- Include ALL work experiences from the entire resume text
- Do not summarize or shorten descriptions - capture full details
- Extract all bullet points as separate achievements
- For multi-page resumes, ensure you capture experiences from ALL pages
- If a project has a URL, include it
- If a project has technologies listed, include them

Resume text to parse:
{text_to_parse}

Return valid JSON only, no markdown formatting:"""
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are an expert resume parser. Extract all information accurately."},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"},
                temperature=0.1,
                max_tokens=12000,  # Increased to handle large resumes with many experiences
            )
            
            import json
            result = json.loads(response.choices[0].message.content)
            
            # Ensure all expected keys exist with proper defaults
            defaults = {
                "full_name": "",
                "email": "",
                "phone": "",
                "location": "",
                "linkedin_url": "",
                "portfolio_url": "",
                "github_url": "",
                "twitter_url": "",
                "medium_url": "",
                "dribbble_url": "",
                "other_url": "",
                "summary": "",
                "skills": [],
                "work_experience": [],
                "education": [],
                "certifications": [],
                "projects": [],
                "languages": [],
            }
            
            # Merge with defaults, ensuring arrays are actually arrays
            final_result = {}
            for key, default_val in defaults.items():
                val = result.get(key, default_val)
                if isinstance(default_val, list) and not isinstance(val, list):
                    val = []
                final_result[key] = val
            
            return final_result
            
        except Exception as e:
            print(f"OpenAI API error in extraction: {e}")
            # Return empty dict to trigger fallback
            return {}
    
    def _enhanced_rule_based_extraction(self, text: str) -> Dict[str, Any]:
        """Enhanced rule-based extraction for all resume sections."""
        lines = text.split('\n')
        
        result = {
            "full_name": "",
            "email": "",
            "phone": "",
            "location": "",
            "linkedin_url": "",
            "portfolio_url": "",
            "summary": "",
            "skills": [],
            "work_experience": [],
            "education": [],
            "certifications": [],
            "projects": [],
            "languages": [],
        }
        
        text_lower = text.lower()
        
        # Extract contact info
        result["email"] = self._extract_email(text) or ""
        result["phone"] = self._extract_phone(text) or ""
        result["linkedin_url"] = self._extract_linkedin(text) or ""
        result["portfolio_url"] = self._extract_portfolio(text) or ""
        result["github_url"] = self._extract_github(text) or ""
        result["twitter_url"] = self._extract_twitter(text) or ""
        result["medium_url"] = self._extract_medium(text) or ""
        result["dribbble_url"] = self._extract_dribbble(text) or ""
        result["other_url"] = self._extract_other_url(text) or ""
        result["location"] = self._extract_location(text) or ""
        result["full_name"] = self._extract_name(lines) or ""
        
        # Extract sections
        result["summary"] = self._extract_summary(text, lines) or ""
        result["skills"] = self._extract_all_skills(text)
        result["work_experience"] = self._extract_work_experience(text, lines)
        result["education"] = self._extract_education(text, lines)
        result["certifications"] = self._extract_certifications(text, lines)
        result["projects"] = self._extract_projects(text, lines)
        result["languages"] = self._extract_languages(text, lines)
        
        return result
    
    def _extract_email(self, text: str) -> Optional[str]:
        pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        match = re.search(pattern, text)
        return match.group(0) if match else None
    
    def _extract_phone(self, text: str) -> Optional[str]:
        patterns = [
            r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\+?\d{1,3}[-.\s]?\d{10}',
            r'\d{3}[-.\s]\d{3}[-.\s]\d{4}',
        ]
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0)
        return None
    
    def _extract_linkedin(self, text: str) -> Optional[str]:
        patterns = [
            r'linkedin\.com/in/[\w-]+',
            r'linkedin\.com/[\w-]+',
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                url = match.group(0)
                if not url.startswith('http'):
                    url = f"https://www.{url}"
                return url
        return None
    
    def _extract_portfolio(self, text: str) -> Optional[str]:
        pattern = r'https?://(?:www\.)?[\w-]+\.[\w-]+[^\s<>\(\)\[\]]*'
        matches = re.findall(pattern, text)
        for url in matches:
            url_lower = url.lower()
            if 'linkedin' not in url_lower and 'github' not in url_lower and 'twitter' not in url_lower and 'medium' not in url_lower and 'dribbble' not in url_lower and 'x.com' not in url_lower:
                return url
        return None
    
    def _extract_github(self, text: str) -> Optional[str]:
        patterns = [
            r'github\.com/[\w-]+',
            r'github\.com/[\w-]+/[\w-]+',
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                url = match.group(0)
                if not url.startswith('http'):
                    url = f"https://{url}"
                return url
        return None
    
    def _extract_twitter(self, text: str) -> Optional[str]:
        patterns = [
            r'(?:twitter|x)\.com/[\w-]+',
            r'@([\w_]+)',
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                url = match.group(0)
                if '@' in url:
                    # It's a handle format
                    handle = match.group(1)
                    return f"https://twitter.com/{handle}"
                if not url.startswith('http'):
                    url = f"https://{url}"
                return url
        return None
    
    def _extract_medium(self, text: str) -> Optional[str]:
        pattern = r'medium\.com/@[\w-]+'
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            url = match.group(0)
            if not url.startswith('http'):
                url = f"https://{url}"
            return url
        return None
    
    def _extract_dribbble(self, text: str) -> Optional[str]:
        pattern = r'dribbble\.com/[\w-]+'
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            url = match.group(0)
            if not url.startswith('http'):
                url = f"https://{url}"
            return url
        return None
    
    def _extract_other_url(self, text: str) -> Optional[str]:
        pattern = r'https?://(?:www\.)?[\w-]+\.[\w-]+[^\s<>\(\)\[\]]*'
        matches = re.findall(pattern, text)
        for url in matches:
            url_lower = url.lower()
            # Skip common known URLs
            skip_domains = ['linkedin', 'github', 'twitter', 'x.com', 'medium', 'dribbble']
            if not any(domain in url_lower for domain in skip_domains):
                return url
        return None
    
    def _extract_location(self, text: str) -> Optional[str]:
        # Common location patterns
        patterns = [
            r'\b([A-Za-z\s]+,\s*(?:AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|ID|IL|IN|IA|KS|KY|LA|ME|MD|MA|MI|MN|MS|MO|MT|NE|NV|NH|NJ|NM|NY|NC|ND|OH|OK|OR|PA|RI|SC|SD|TN|TX|UT|VT|VA|WA|WV|WI|WY))\b',
            r'\b([A-Za-z\s]+,\s*(?:India|USA|UK|Canada|Australia|Germany|France))\b',
        ]
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1)
        return None
    
    def _extract_name(self, lines: List[str]) -> Optional[str]:
        for line in lines[:10]:
            line = line.strip()
            if line and len(line) > 2 and len(line) < 50:
                if re.match(r'^[A-Za-z\s\.]+$', line) and len(line.split()) <= 4:
                    if not any(word in line.lower() for word in ['resume', 'cv', 'curriculum', 'page', 'email', 'phone']):
                        return line
        return None
    
    def _extract_summary(self, text: str, lines: List[str]) -> Optional[str]:
        # Look for summary/professional summary section
        summary_keywords = ['summary', 'professional summary', 'career objective', 'objective', 'about me', 'profile']
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in summary_keywords):
                # Next non-empty line is likely the summary
                for j in range(i+1, min(i+10, len(lines))):
                    summary = lines[j].strip()
                    if len(summary) > 50:
                        return summary
        
        # Fallback: find first long paragraph
        for line in lines:
            line = line.strip()
            if len(line) > 100 and len(line) < 1000:
                if not any(x in line.lower() for x in ['experience', 'education', 'skills']):
                    return line
        return None
    
    def _extract_all_skills(self, text: str) -> List[str]:
        """Extract comprehensive skills list."""
        skill_keywords = [
            # Programming
            'python', 'javascript', 'js', 'typescript', 'ts', 'java', 'c++', 'c#', 'csharp', 'go', 'golang',
            'rust', 'ruby', 'php', 'swift', 'kotlin', 'scala', 'r', 'matlab', 'perl', 'lua', 'dart',
            # Web
            'html', 'css', 'sass', 'scss', 'less', 'bootstrap', 'tailwind', 'react', 'angular', 'vue',
            'next.js', 'nuxt', 'gatsby', 'svelte', 'jquery', 'ajax', 'json', 'xml', 'yaml',
            # Backend
            'node.js', 'nodejs', 'express', 'django', 'flask', 'fastapi', 'spring', 'spring boot',
            'laravel', 'rails', 'asp.net', 'graphql', 'rest api', 'soap', 'grpc',
            # Databases
            'sql', 'mysql', 'postgresql', 'postgres', 'mongodb', 'mongoose', 'redis', 'sqlite',
            'dynamodb', 'cassandra', 'elasticsearch', 'firebase', 'supabase', 'prisma', 'sequelize',
            # Cloud & DevOps
            'aws', 'amazon web services', 'azure', 'gcp', 'google cloud', 'heroku', 'netlify', 'vercel',
            'docker', 'kubernetes', 'k8s', 'jenkins', 'gitlab ci', 'github actions', 'circleci',
            'terraform', 'ansible', 'puppet', 'chef', 'nginx', 'apache', 'linux', 'unix', 'bash',
            # Data Science
            'machine learning', 'ml', 'deep learning', 'tensorflow', 'pytorch', 'keras', 'scikit-learn',
            'pandas', 'numpy', 'matplotlib', 'seaborn', 'plotly', 'jupyter', 'rstudio', 'spss',
            'data analysis', 'data visualization', 'statistics', 'nlp', 'computer vision', 'opencv',
            # Mobile
            'react native', 'flutter', 'ios', 'android', 'xamarin', 'ionic', 'cordova', 'phonegap',
            # Testing
            'jest', 'mocha', 'chai', 'cypress', 'selenium', 'junit', 'pytest', 'cucumber', 'cicd',
            # Tools
            'git', 'github', 'gitlab', 'bitbucket', 'jira', 'confluence', 'trello', 'asana',
            'slack', 'teams', 'zoom', 'postman', 'insomnia', 'figma', 'sketch', 'adobe xd',
            # Methodologies
            'agile', 'scrum', 'kanban', 'waterfall', 'devops', 'ci/cd', 'tdd', 'bdd', 'oop',
            # Soft skills
            'leadership', 'communication', 'teamwork', 'problem solving', 'critical thinking',
            'time management', 'project management', 'collaboration', 'mentoring',
        ]
        
        text_lower = text.lower()
        found_skills = []
        
        for skill in skill_keywords:
            # Look for word boundaries
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                # Normalize skill name
                normalized = skill.title() if ' ' in skill else skill.upper() if len(skill) <= 3 else skill.title()
                found_skills.append(normalized)
        
        return list(set(found_skills))
    
    def _extract_work_experience(self, text: str, lines: List[str]) -> List[Dict]:
        """Extract work experience entries."""
        experiences = []
        
        # Find experience section
        exp_section_start = -1
        exp_keywords = ['experience', 'work experience', 'professional experience', 'employment', 'career history']
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in exp_keywords):
                exp_section_start = i
                break
        
        if exp_section_start == -1:
            return experiences
        
        # Look for education or other sections to know where to stop
        section_end = len(lines)
        next_sections = ['education', 'skills', 'projects', 'certifications', 'languages', 'interests', 'references']
        
        for i in range(exp_section_start + 1, len(lines)):
            line_lower = lines[i].lower().strip()
            if any(keyword in line_lower for keyword in next_sections):
                section_end = i
                break
        
        # Parse experience entries
        exp_text = '\n'.join(lines[exp_section_start:section_end])
        
        # Try to identify job entries (title + company pattern)
        # Common patterns: "Job Title, Company" or "Job Title | Company" or "Job Title at Company"
        job_patterns = [
            r'([A-Za-z\s]+(?:Engineer|Developer|Manager|Director|Analyst|Designer|Consultant|Specialist|Lead|Architect|Intern|Trainee)[A-Za-z\s]*)[,\|]\s*([^\n]+)',
            r'([A-Za-z\s]+)at\s+([A-Za-z\s]+)',
        ]
        
        # Extract entries based on date patterns
        date_pattern = r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{1,2}/\d{4}|\d{4})\s*[-–]\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{1,2}/\d{4}|\d{4}|Present|Current)'
        
        dates = list(re.finditer(date_pattern, exp_text, re.IGNORECASE))
        
        for i, date_match in enumerate(dates):
            start_date = date_match.group(1)
            end_date = date_match.group(2)
            is_current = end_date.lower() in ['present', 'current']
            
            # Look for job title and company before this date
            context_before = exp_text[max(0, date_match.start() - 500):date_match.start()]
            
            title = "Unknown Position"
            company = "Unknown Company"
            location = ""
            
            # Try to extract title and company
            for pattern in job_patterns:
                match = re.search(pattern, context_before)
                if match:
                    title = match.group(1).strip()
                    company = match.group(2).strip()
                    break
            
            # Look for description after the date
            context_after = exp_text[date_match.end():min(len(exp_text), date_match.end() + 800)]
            next_date = dates[i + 1] if i + 1 < len(dates) else None
            if next_date:
                context_after = exp_text[date_match.end():max(0, next_date.start() - len(exp_text[:exp_section_start]))]
            
            # Extract description (bullet points or paragraph)
            description_lines = []
            achievements = []
            
            for line in context_after.split('\n')[:15]:
                line = line.strip()
                if line and len(line) > 10:
                    if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                        achievements.append(line.lstrip('•-* ').strip())
                    elif len(line) > 30:
                        description_lines.append(line)
            
            description = ' '.join(description_lines[:3]) if description_lines else ""
            
            exp = {
                "title": title,
                "company": company,
                "location": location,
                "start_date": start_date,
                "end_date": "" if is_current else end_date,
                "is_current": is_current,
                "description": description,
                "achievements": achievements[:5],
            }
            experiences.append(exp)
        
        return experiences[:10]  # Limit to 10 most recent
    
    def _extract_education(self, text: str, lines: List[str]) -> List[Dict]:
        """Extract education entries."""
        education = []
        
        edu_section_start = -1
        edu_keywords = ['education', 'academic background', 'academic qualifications', 'qualifications']
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in edu_keywords):
                edu_section_start = i
                break
        
        if edu_section_start == -1:
            return education
        
        # Find where section ends
        section_end = len(lines)
        next_sections = ['experience', 'skills', 'projects', 'certifications', 'languages']
        
        for i in range(edu_section_start + 1, len(lines)):
            line_lower = lines[i].lower().strip()
            if any(keyword in line_lower for keyword in next_sections):
                section_end = i
                break
        
        edu_text = '\n'.join(lines[edu_section_start:section_end])
        
        # Degree patterns
        degree_patterns = [
            r'(Bachelor[\'s]*\s+(?:of\s+)?(?:Science|Arts|Engineering|Technology|Business Administration|B\.S\.|B\.A\.|B\.E\.|B\.Tech|BE|BTech|BS|BA))',
            r'(Master[\'s]*\s+(?:of\s+)?(?:Science|Arts|Engineering|Business Administration|Technology|M\.S\.|M\.A\.|M\.E\.|M\.Tech|MBA|MS|MA|ME|MTech))',
            r'(Ph\.?D\.?|Doctor of Philosophy)',
            r'(Associate[\'s]*\s+(?:of\s+)?(?:Arts|Science|A\.A\.|A\.S\.|AA|AS))',
            r'(High School|Diploma|Certificate|Certification)',
        ]
        
        degrees_found = []
        for pattern in degree_patterns:
            for match in re.finditer(pattern, edu_text, re.IGNORECASE):
                degrees_found.append((match.start(), match.group(1)))
        
        # Sort by position in text
        degrees_found.sort(key=lambda x: x[0])
        
        for i, (pos, degree) in enumerate(degrees_found):
            # Look for university name near this degree
            context = edu_text[max(0, pos - 200):pos + 300]
            
            institution = "Unknown Institution"
            institution_patterns = [
                r'(?:University|College|Institute|School)\s+(?:of\s+)?[A-Za-z\s]+',
                r'[A-Za-z\s]+(?:University|College|Institute|School)',
            ]
            for pattern in institution_patterns:
                match = re.search(pattern, context)
                if match:
                    institution = match.group(0).strip()
                    break
            
            # Look for graduation year
            year_match = re.search(r'\b(19|20)\d{2}\b', context)
            graduation_date = year_match.group(1) if year_match else ""
            
            # Look for GPA
            gpa_match = re.search(r'GPA[:\s]*([\d.]+)', context, re.IGNORECASE)
            gpa = gpa_match.group(1) if gpa_match else ""
            
            edu = {
                "degree": degree,
                "institution": institution,
                "location": "",
                "graduation_date": graduation_date,
                "gpa": gpa,
            }
            education.append(edu)
        
        return education
    
    def _extract_certifications(self, text: str, lines: List[str]) -> List[Dict]:
        """Extract certifications."""
        certs = []
        
        cert_section_start = -1
        cert_keywords = ['certifications', 'certificates', 'licenses', 'professional certifications']
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in cert_keywords):
                cert_section_start = i
                break
        
        if cert_section_start == -1:
            return certs
        
        section_end = len(lines)
        next_sections = ['experience', 'education', 'skills', 'projects', 'languages']
        
        for i in range(cert_section_start + 1, len(lines)):
            line_lower = lines[i].lower().strip()
            if any(keyword in line_lower for keyword in next_sections):
                section_end = i
                break
        
        cert_text = '\n'.join(lines[cert_section_start:section_end])
        
        # Look for certification patterns
        cert_patterns = [
            r'(?:AWS|Amazon Web Services)[\s\w]*Certified[\w\s]*',
            r'(?:Microsoft|Google|Oracle|Cisco|CompTIA|PMP|CISSP|CEH)[\s\w]*Certified[\w\s]*',
            r'Certified[\s\w]+Professional',
        ]
        
        for pattern in cert_patterns:
            for match in re.finditer(pattern, cert_text, re.IGNORECASE):
                cert = {
                    "name": match.group(0).strip(),
                    "issuer": "Unknown",
                    "date": "",
                }
                certs.append(cert)
        
        return certs[:10]
    
    def _extract_projects(self, text: str, lines: List[str]) -> List[Dict]:
        """Extract project entries."""
        projects = []
        
        proj_section_start = -1
        proj_keywords = ['projects', 'personal projects', 'academic projects', 'side projects']
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in proj_keywords):
                proj_section_start = i
                break
        
        if proj_section_start == -1:
            return projects
        
        section_end = len(lines)
        next_sections = ['experience', 'education', 'skills', 'certifications', 'languages']
        
        for i in range(proj_section_start + 1, len(lines)):
            line_lower = lines[i].lower().strip()
            if any(keyword in line_lower for keyword in next_sections):
                section_end = i
                break
        
        proj_text = '\n'.join(lines[proj_section_start:section_end])
        proj_lines = [l.strip() for l in proj_text.split('\n') if l.strip()]
        
        # Simple project extraction - lines that look like project names
        for line in proj_lines[1:]:  # Skip header
            if len(line) > 10 and len(line) < 100:
                if not any(x in line.lower() for x in ['http', 'github', 'description']):
                    project = {
                        "name": line,
                        "description": "",
                        "technologies": [],
                        "url": "",
                    }
                    projects.append(project)
        
        return projects[:5]
    
    def _extract_languages(self, text: str, lines: List[str]) -> List[Dict]:
        """Extract languages."""
        languages = []
        
        lang_section_start = -1
        lang_keywords = ['languages', 'language proficiency', 'spoken languages']
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in lang_keywords):
                lang_section_start = i
                break
        
        if lang_section_start == -1:
            return languages
        
        section_end = len(lines)
        next_sections = ['experience', 'education', 'skills', 'certifications', 'projects']
        
        for i in range(lang_section_start + 1, len(lines)):
            line_lower = lines[i].lower().strip()
            if any(keyword in line_lower for keyword in next_sections):
                section_end = i
                break
        
        lang_text = '\n'.join(lines[lang_section_start:section_end])
        
        # Common languages
        common_languages = [
            'english', 'spanish', 'french', 'german', 'chinese', 'japanese', 'korean',
            'arabic', 'hindi', 'portuguese', 'russian', 'italian', 'dutch', 'turkish',
            'tamil', 'telugu', 'marathi', 'bengali', 'urdu', 'gujarati', 'kannada',
        ]
        
        text_lower = lang_text.lower()
        
        for lang in common_languages:
            if lang in text_lower:
                proficiency = "Conversational"
                # Try to find proficiency level
                patterns = [
                    rf'{lang}[\s\w]*native',
                    rf'{lang}[\s\w]*fluent',
                    rf'{lang}[\s\w]*professional',
                    rf'{lang}[\s\w]*intermediate',
                    rf'{lang}[\s\w]*basic',
                ]
                levels = ["Native", "Fluent", "Professional", "Intermediate", "Basic"]
                
                for pattern, level in zip(patterns, levels):
                    if re.search(pattern, text_lower):
                        proficiency = level
                        break
                
                languages.append({
                    "language": lang.title(),
                    "proficiency": proficiency,
                })
        
        return languages
    
    def _calculate_comprehensive_ats_score(self, parsed_data: Dict[str, Any]) -> int:
        """Calculate comprehensive ATS score."""
        score = 0
        max_score = 100
        
        # Contact Information (15 points)
        if parsed_data.get("email"):
            score += 5
        if parsed_data.get("phone"):
            score += 5
        if parsed_data.get("location"):
            score += 5
        
        # Professional Summary (10 points)
        if parsed_data.get("summary") and len(parsed_data["summary"]) > 50:
            score += 10
        
        # Skills (20 points)
        skills = parsed_data.get("skills", [])
        if len(skills) >= 15:
            score += 20
        elif len(skills) >= 10:
            score += 15
        elif len(skills) >= 5:
            score += 10
        else:
            score += 5
        
        # Work Experience (25 points)
        exp = parsed_data.get("work_experience", [])
        if len(exp) >= 3:
            score += 25
        elif len(exp) == 2:
            score += 20
        elif len(exp) == 1:
            score += 15
        
        # Education (10 points)
        edu = parsed_data.get("education", [])
        if len(edu) >= 2:
            score += 10
        elif len(edu) == 1:
            score += 8
        
        # Certifications (5 points)
        if parsed_data.get("certifications"):
            score += 5
        
        # Projects (5 points)
        if parsed_data.get("projects"):
            score += 5
        
        # LinkedIn/Portfolio/Social Links (5 points)
        if any([parsed_data.get("linkedin_url"), parsed_data.get("portfolio_url"), 
                parsed_data.get("github_url"), parsed_data.get("twitter_url")]):
            score += 5
        
        # Languages (5 points)
        if parsed_data.get("languages"):
            score += 5
        
        return min(score, max_score)


# Singleton instance
resume_parser = ResumeParser()
