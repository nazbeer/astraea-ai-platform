"""
ATS-Friendly Resume Generator
Generates professional resumes in PDF and DOCX formats optimized for Applicant Tracking Systems.
Supports 3 different ATS-friendly templates.
"""

from typing import Dict, List, Any, Optional
from datetime import datetime
import os
import io

# PDF Generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListFlowable, ListItem, HRFlowable
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

# DOCX Generation
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


class ResumeGenerator:
    """Generate ATS-friendly resumes in PDF and DOCX formats with multiple templates."""
    
    # Template names
    TEMPLATES = {
        'modern': 'Modern ATS-Friendly',
        'classic': 'Classic Professional', 
        'minimal': 'Minimal Clean'
    }
    
    def __init__(self, output_dir: str = "storage/resumes"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def generate_pdf(self, resume_data: Dict[str, Any], filename: Optional[str] = None, template: str = 'modern', return_bytes: bool = False) -> str | bytes:
        """Generate an ATS-friendly PDF resume with selected template."""
        if filename is None:
            filename = f"resume_{resume_data.get('user_id', 'unknown')}_{int(datetime.now().timestamp())}.pdf"
        
        filepath = os.path.join(self.output_dir, filename)
        
        # Use BytesIO if returning bytes
        if return_bytes:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=0.75*inch,
                leftMargin=0.75*inch,
                topMargin=0.75*inch,
                bottomMargin=0.75*inch
            )
        else:
            doc = SimpleDocTemplate(
                filepath,
                pagesize=letter,
                rightMargin=0.75*inch,
                leftMargin=0.75*inch,
                topMargin=0.75*inch,
                bottomMargin=0.75*inch
            )
        
        # Route to appropriate template
        if template == 'classic':
            story = self._build_classic_template(resume_data)
        elif template == 'minimal':
            story = self._build_minimal_template(resume_data)
        else:  # default modern
            story = self._build_modern_template(resume_data)
        
        doc.build(story)
        
        if return_bytes:
            buffer.seek(0)
            return buffer.getvalue()
        
        return filepath
    
    def _build_modern_template(self, resume_data: Dict[str, Any]) -> List:
        """Build modern ATS-friendly template with clean sections."""
        styles = getSampleStyleSheet()
        
        # Modern styles
        name_style = ParagraphStyle(
            'ModernName',
            parent=styles['Heading1'],
            fontSize=22,
            spaceAfter=6,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#1a1a1a'),
            fontName='Helvetica-Bold'
        )
        
        contact_style = ParagraphStyle(
            'ModernContact',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=2,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#4a4a4a'),
            fontName='Helvetica'
        )
        
        section_style = ParagraphStyle(
            'ModernSection',
            parent=styles['Heading2'],
            fontSize=12,
            spaceAfter=8,
            spaceBefore=14,
            textColor=colors.HexColor('#1a1a1a'),
            fontName='Helvetica-Bold',
            borderColor=colors.HexColor('#1a1a1a'),
            borderWidth=2,
            borderPadding=4
        )
        
        job_title_style = ParagraphStyle(
            'ModernJobTitle',
            parent=styles['Normal'],
            fontSize=11,
            fontName='Helvetica-Bold',
            spaceAfter=2,
            textColor=colors.HexColor('#1a1a1a')
        )
        
        company_style = ParagraphStyle(
            'ModernCompany',
            parent=styles['Normal'],
            fontSize=10,
            fontName='Helvetica-Bold',
            spaceAfter=1,
            textColor=colors.HexColor('#333333')
        )
        
        date_style = ParagraphStyle(
            'ModernDate',
            parent=styles['Normal'],
            fontSize=9,
            fontName='Helvetica-Oblique',
            textColor=colors.HexColor('#666666'),
            spaceAfter=3
        )
        
        normal_style = ParagraphStyle(
            'ModernNormal',
            parent=styles['Normal'],
            fontSize=10,
            fontName='Helvetica',
            spaceAfter=3,
            leading=14,
            textColor=colors.HexColor('#333333')
        )
        
        return self._build_resume_story(resume_data, name_style, contact_style, section_style, 
                                        job_title_style, company_style, date_style, normal_style)
    
    def _build_classic_template(self, resume_data: Dict[str, Any]) -> List:
        """Build classic professional template."""
        styles = getSampleStyleSheet()
        
        # Classic styles - more traditional
        name_style = ParagraphStyle(
            'ClassicName',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=8,
            alignment=TA_CENTER,
            textColor=colors.black,
            fontName='Times-Bold'
        )
        
        contact_style = ParagraphStyle(
            'ClassicContact',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=2,
            alignment=TA_CENTER,
            textColor=colors.black,
            fontName='Times-Roman'
        )
        
        section_style = ParagraphStyle(
            'ClassicSection',
            parent=styles['Heading2'],
            fontSize=12,
            spaceAfter=6,
            spaceBefore=12,
            textColor=colors.black,
            fontName='Times-Bold',
            underline=True
        )
        
        job_title_style = ParagraphStyle(
            'ClassicJobTitle',
            parent=styles['Normal'],
            fontSize=11,
            fontName='Times-Bold',
            spaceAfter=1
        )
        
        company_style = ParagraphStyle(
            'ClassicCompany',
            parent=styles['Normal'],
            fontSize=10,
            fontName='Times-Roman',
            spaceAfter=1
        )
        
        date_style = ParagraphStyle(
            'ClassicDate',
            parent=styles['Normal'],
            fontSize=9,
            fontName='Times-Italic',
            textColor=colors.black,
            spaceAfter=2
        )
        
        normal_style = ParagraphStyle(
            'ClassicNormal',
            parent=styles['Normal'],
            fontSize=10,
            fontName='Times-Roman',
            spaceAfter=2,
            leading=13
        )
        
        return self._build_resume_story(resume_data, name_style, contact_style, section_style,
                                        job_title_style, company_style, date_style, normal_style)
    
    def _build_minimal_template(self, resume_data: Dict[str, Any]) -> List:
        """Build minimal clean template."""
        styles = getSampleStyleSheet()
        
        # Minimal styles - very clean and simple
        name_style = ParagraphStyle(
            'MinimalName',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=4,
            alignment=TA_LEFT,
            textColor=colors.black,
            fontName='Helvetica-Bold'
        )
        
        contact_style = ParagraphStyle(
            'MinimalContact',
            parent=styles['Normal'],
            fontSize=9,
            spaceAfter=1,
            alignment=TA_LEFT,
            textColor=colors.HexColor('#555555'),
            fontName='Helvetica'
        )
        
        section_style = ParagraphStyle(
            'MinimalSection',
            parent=styles['Heading2'],
            fontSize=10,
            spaceAfter=4,
            spaceBefore=10,
            textColor=colors.HexColor('#555555'),
            fontName='Helvetica-Bold',
            underline=True
        )
        
        job_title_style = ParagraphStyle(
            'MinimalJobTitle',
            parent=styles['Normal'],
            fontSize=10,
            fontName='Helvetica-Bold',
            spaceAfter=0
        )
        
        company_style = ParagraphStyle(
            'MinimalCompany',
            parent=styles['Normal'],
            fontSize=9,
            fontName='Helvetica',
            spaceAfter=0
        )
        
        date_style = ParagraphStyle(
            'MinimalDate',
            parent=styles['Normal'],
            fontSize=8,
            fontName='Helvetica',
            textColor=colors.HexColor('#777777'),
            spaceAfter=2
        )
        
        normal_style = ParagraphStyle(
            'MinimalNormal',
            parent=styles['Normal'],
            fontSize=9,
            fontName='Helvetica',
            spaceAfter=2,
            leading=12
        )
        
        return self._build_resume_story(resume_data, name_style, contact_style, section_style,
                                        job_title_style, company_style, date_style, normal_style)
    
    def _build_resume_story(self, resume_data: Dict[str, Any], name_style, contact_style, 
                           section_style, job_title_style, company_style, date_style, normal_style) -> List:
        """Build the resume content with given styles."""
        story = []
        
        # Header - Name
        full_name = resume_data.get('full_name', '')
        story.append(Paragraph(full_name, name_style))
        
        # Contact Information
        contact_parts = []
        if resume_data.get('location'):
            contact_parts.append(resume_data['location'])
        if resume_data.get('email'):
            contact_parts.append(resume_data['email'])
        if resume_data.get('phone'):
            contact_parts.append(resume_data['phone'])
        if resume_data.get('linkedin_url'):
            linkedin = resume_data['linkedin_url'].replace('https://', '').replace('http://', '')
            contact_parts.append(f"linkedin.com/in/{linkedin.split('/')[-1]}")
        if resume_data.get('portfolio_url'):
            portfolio = resume_data['portfolio_url'].replace('https://', '').replace('http://', '')
            contact_parts.append(portfolio)
        
        if contact_parts:
            story.append(Paragraph(" | ".join(contact_parts), contact_style))
        story.append(Spacer(1, 12))
        
        # Summary
        if resume_data.get('summary'):
            story.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
            story.append(Paragraph(resume_data['summary'], normal_style))
            story.append(Spacer(1, 6))
        
        # Skills
        if resume_data.get('skills'):
            story.append(Paragraph("SKILLS", section_style))
            skills_text = ", ".join(resume_data['skills'])
            story.append(Paragraph(skills_text, normal_style))
            story.append(Spacer(1, 6))
        
        # Work Experience
        if resume_data.get('work_experience'):
            story.append(Paragraph("WORK EXPERIENCE", section_style))
            for exp in resume_data['work_experience']:
                # Job title and company on same line
                title_text = f"{exp.get('title', '')}"
                story.append(Paragraph(title_text, job_title_style))
                
                company_text = f"{exp.get('company', '')}"
                if exp.get('location'):
                    company_text += f" | {exp['location']}"
                story.append(Paragraph(company_text, company_style))
                
                # Dates
                date_text = f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present') if not exp.get('is_current') else 'Present'}"
                story.append(Paragraph(date_text, date_style))
                
                # Description
                if exp.get('description'):
                    story.append(Paragraph(exp['description'], normal_style))
                
                # Achievements
                for achievement in exp.get('achievements', []):
                    if achievement:
                        story.append(Paragraph(f"• {achievement}", normal_style))
                
                story.append(Spacer(1, 8))
        
        # Education
        if resume_data.get('education'):
            story.append(Paragraph("EDUCATION", section_style))
            for edu in resume_data['education']:
                degree_text = f"{edu.get('degree', '')}"
                story.append(Paragraph(degree_text, job_title_style))
                
                school_text = f"{edu.get('institution', '')}"
                if edu.get('location'):
                    school_text += f" | {edu['location']}"
                story.append(Paragraph(school_text, company_style))
                
                grad_text = f"{edu.get('graduation_date', '')}"
                if edu.get('gpa'):
                    grad_text += f" | GPA: {edu['gpa']}"
                story.append(Paragraph(grad_text, date_style))
                story.append(Spacer(1, 6))
        
        # Certifications
        if resume_data.get('certifications'):
            story.append(Paragraph("CERTIFICATIONS", section_style))
            for cert in resume_data['certifications']:
                cert_text = f"{cert.get('name', '')}"
                if cert.get('issuer'):
                    cert_text += f" - {cert['issuer']}"
                if cert.get('date'):
                    cert_text += f" ({cert['date']})"
                story.append(Paragraph(f"• {cert_text}", normal_style))
            story.append(Spacer(1, 6))
        
        # Projects
        if resume_data.get('projects'):
            story.append(Paragraph("PROJECTS", section_style))
            for proj in resume_data['projects']:
                proj_text = f"{proj.get('name', '')}"
                story.append(Paragraph(proj_text, job_title_style))
                
                if proj.get('description'):
                    story.append(Paragraph(proj['description'], normal_style))
                
                if proj.get('technologies'):
                    tech_text = f"Technologies: {', '.join(proj['technologies'])}"
                    story.append(Paragraph(tech_text, date_style))
                
                if proj.get('url'):
                    story.append(Paragraph(f"URL: {proj['url']}", date_style))
                
                story.append(Spacer(1, 6))
        
        # Languages
        if resume_data.get('languages'):
            story.append(Paragraph("LANGUAGES", section_style))
            for lang in resume_data['languages']:
                lang_text = f"{lang.get('language', '')}"
                if lang.get('proficiency'):
                    lang_text += f" - {lang['proficiency']}"
                story.append(Paragraph(f"• {lang_text}", normal_style))
        
        return story
    
    def generate_docx(self, resume_data: Dict[str, Any], filename: Optional[str] = None, template: str = 'modern', return_bytes: bool = False) -> str | bytes:
        """Generate an ATS-friendly DOCX resume."""
        if filename is None:
            filename = f"resume_{resume_data.get('user_id', 'unknown')}_{int(datetime.now().timestamp())}.docx"
        
        filepath = os.path.join(self.output_dir, filename)
        
        doc = Document()
        
        # Set narrow margins for more content
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Header - Name
        name = doc.add_heading(resume_data.get('full_name', ''), level=0)
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact Information
        contact_parts = []
        if resume_data.get('location'):
            contact_parts.append(resume_data['location'])
        if resume_data.get('email'):
            contact_parts.append(resume_data['email'])
        if resume_data.get('phone'):
            contact_parts.append(resume_data['phone'])
        if resume_data.get('linkedin_url'):
            linkedin = resume_data['linkedin_url'].replace('https://', '').replace('http://', '')
            contact_parts.append(f"linkedin.com/in/{linkedin.split('/')[-1]}")
        
        if contact_parts:
            contact_para = doc.add_paragraph(" | ".join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary
        if resume_data.get('summary'):
            doc.add_heading('PROFESSIONAL SUMMARY', level=1)
            doc.add_paragraph(resume_data['summary'])
        
        # Skills
        if resume_data.get('skills'):
            doc.add_heading('SKILLS', level=1)
            doc.add_paragraph(", ".join(resume_data['skills']))
        
        # Work Experience
        if resume_data.get('work_experience'):
            doc.add_heading('WORK EXPERIENCE', level=1)
            for exp in resume_data['work_experience']:
                # Job Title
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(exp.get('title', ''))
                title_run.bold = True
                title_run.font.size = Pt(11)
                
                # Company and Location
                company_text = f"{exp.get('company', '')}"
                if exp.get('location'):
                    company_text += f" | {exp['location']}"
                company_para = doc.add_paragraph(company_text)
                company_para.runs[0].italic = True
                
                # Dates
                date_text = f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present') if not exp.get('is_current') else 'Present'}"
                date_para = doc.add_paragraph(date_text)
                date_para.runs[0].italic = True
                date_para.runs[0].font.size = Pt(9)
                
                # Description
                if exp.get('description'):
                    doc.add_paragraph(exp['description'])
                
                # Achievements
                for achievement in exp.get('achievements', []):
                    if achievement:
                        doc.add_paragraph(f"• {achievement}")
                
                doc.add_paragraph()
        
        # Education
        if resume_data.get('education'):
            doc.add_heading('EDUCATION', level=1)
            for edu in resume_data['education']:
                # Degree
                degree_para = doc.add_paragraph()
                degree_run = degree_para.add_run(edu.get('degree', ''))
                degree_run.bold = True
                degree_run.font.size = Pt(11)
                
                # School and Location
                school_text = f"{edu.get('institution', '')}"
                if edu.get('location'):
                    school_text += f" | {edu['location']}"
                doc.add_paragraph(school_text)
                
                # Graduation Date
                grad_text = f"{edu.get('graduation_date', '')}"
                if edu.get('gpa'):
                    grad_text += f" | GPA: {edu['gpa']}"
                grad_para = doc.add_paragraph(grad_text)
                grad_para.runs[0].italic = True
                grad_para.runs[0].font.size = Pt(9)
        
        # Certifications
        if resume_data.get('certifications'):
            doc.add_heading('CERTIFICATIONS', level=1)
            for cert in resume_data['certifications']:
                cert_text = f"{cert.get('name', '')}"
                if cert.get('issuer'):
                    cert_text += f" - {cert['issuer']}"
                if cert.get('date'):
                    cert_text += f" ({cert['date']})"
                doc.add_paragraph(f"• {cert_text}")
        
        # Projects
        if resume_data.get('projects'):
            doc.add_heading('PROJECTS', level=1)
            for proj in resume_data['projects']:
                # Project Name
                proj_para = doc.add_paragraph()
                proj_run = proj_para.add_run(proj.get('name', ''))
                proj_run.bold = True
                
                if proj.get('description'):
                    doc.add_paragraph(proj['description'])
                
                if proj.get('technologies'):
                    tech_para = doc.add_paragraph(f"Technologies: {', '.join(proj['technologies'])}")
                    tech_para.runs[0].italic = True
                    tech_para.runs[0].font.size = Pt(9)
                
                if proj.get('url'):
                    doc.add_paragraph(f"URL: {proj['url']}")
        
        # Languages
        if resume_data.get('languages'):
            doc.add_heading('LANGUAGES', level=1)
            for lang in resume_data['languages']:
                lang_text = f"{lang.get('language', '')}"
                if lang.get('proficiency'):
                    lang_text += f" - {lang['proficiency']}"
                doc.add_paragraph(f"• {lang_text}")
        
        if return_bytes:
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        else:
            doc.save(filepath)
            return filepath
    
    def calculate_ats_score(self, resume_data: Dict[str, Any]) -> Dict[str, Any]:
        """Calculate ATS compatibility score and provide suggestions."""
        score = 0
        max_score = 100
        suggestions = []
        keywords = []
        
        # Contact Information (15 points)
        if resume_data.get('email'):
            score += 5
        else:
            suggestions.append("Add an email address")
        
        if resume_data.get('phone'):
            score += 5
        else:
            suggestions.append("Add a phone number")
        
        if resume_data.get('location'):
            score += 5
        else:
            suggestions.append("Add your location")
        
        # Professional Summary (10 points)
        if resume_data.get('summary') and len(resume_data['summary']) > 50:
            score += 10
        else:
            suggestions.append("Add a professional summary (at least 50 characters)")
        
        # Skills (20 points)
        skills = resume_data.get('skills', [])
        if len(skills) >= 15:
            score += 20
        elif len(skills) >= 10:
            score += 15
        elif len(skills) >= 5:
            score += 10
        else:
            score += 5
            suggestions.append("Add more skills (aim for at least 10)")
        
        keywords.extend(skills)
        
        # Work Experience (25 points)
        exp = resume_data.get('work_experience', [])
        if len(exp) >= 3:
            score += 25
        elif len(exp) == 2:
            score += 20
        elif len(exp) == 1:
            score += 15
        else:
            suggestions.append("Add work experience")
        
        # Check for achievements in work experience
        has_achievements = any(
            any(a for a in e.get('achievements', []))
            for e in exp
        )
        if not has_achievements:
            suggestions.append("Add achievements to your work experience")
        
        # Education (10 points)
        edu = resume_data.get('education', [])
        if len(edu) >= 2:
            score += 10
        elif len(edu) == 1:
            score += 8
        else:
            suggestions.append("Add your education")
        
        # Certifications (5 points)
        if resume_data.get('certifications'):
            score += 5
        
        # Projects (5 points)
        if resume_data.get('projects'):
            score += 5
        
        # LinkedIn/Portfolio/Social Links (5 points)
        if any([resume_data.get('linkedin_url'), resume_data.get('portfolio_url'), 
                resume_data.get('github_url'), resume_data.get('twitter_url')]):
            score += 5
        else:
            suggestions.append("Add a LinkedIn profile or portfolio URL")
        
        # Languages (5 points)
        if resume_data.get('languages'):
            score += 5
        
        # Deduct points for common ATS issues
        if resume_data.get('summary') and len(resume_data['summary']) > 500:
            suggestions.append("Summary is too long - keep it under 500 characters")
        
        # Check for keywords in work experience
        for exp_item in exp:
            desc = exp_item.get('description', '')
            if desc:
                # Extract potential keywords
                words = desc.lower().split()
                keywords.extend([w for w in words if len(w) > 4])
        
        # Remove duplicates and limit
        keywords = list(set([k.lower() for k in keywords]))[:20]
        
        return {
            "score": min(score, max_score),
            "suggestions": suggestions[:5],  # Top 5 suggestions
            "keywords": keywords,
            "is_ats_friendly": score >= 70
        }


# Singleton instance
resume_generator = ResumeGenerator()
